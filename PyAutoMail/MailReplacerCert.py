"""
cd PyAutoMail
python MailReplacerCert.py
"""

from AutoMailer import *
import os,shutil
import re
UserDetails = read_user_details()
MainDir = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR =MainDir



import fnmatch
import os
import zipfile
import shutil
from docx import Document

def ReplaceText(InputFilepath, OutputFilepath, ReplacementDictionary):
    document = Document(InputFilepath)
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            print(text)
            if text in ReplacementDictionary.keys():
                text = text.replace(text, ReplacementDictionary[text])
                inline[i].text = text

    document.save(OutputFilepath)

def find(pattern, path):
    result = []
    for root, dirs, files in os.walk(path):
        for name in files:
            if fnmatch.fnmatch(name, pattern):
                result.append(os.path.join(root, name))
    return result

def ExtractZipFile(SourceZipFile,OutputDirectory):
    with zipfile.ZipFile(SourceZipFile, 'r') as zip_ref:
        zip_ref.extractall(OutputDirectory)

def ReplaceWordImage(SrcImgFilePath,DestImgFileName,OutputDirectory):
    DestImgFilePath=os.path.join(OutputDirectory, "word", "media",DestImgFileName)
    # First Delete the Destination File & Then Copy the SourceFile to the DestinationFile
    os.remove(DestImgFilePath)
    shutil.copy(SrcImgFilePath, DestImgFilePath)

def zip_directory(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, mode='w') as zipf:
        len_dir_path = len(folder_path)
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, file_path[len_dir_path:])

def ReconsructDocument(OutputDirectory,OutputFilepath):
    zip_directory(OutputDirectory,OutputFilepath)
    
def create_dir(dir):
  if not os.path.exists(dir):
    os.makedirs(dir)
    # print("Created Directory : ", dir)
    return dir
def readHTML(FilePath):
    with open(FilePath, "r") as f:
        return f.readlines()


def writeHTML(FilePath, txt):
    with open(FilePath, "w") as f:
        for row in txt:
            if not str(row).strip() == "":
                f.write(str(row) + '\n')


def addtoHTMLFile(txt, pos, str_to_add, fpath):
    sub_txt = txt[:pos]
    sub_txt.extend(str_to_add.split('\n'))
    sub_txt.extend(txt[pos:])
    writeHTML(fpath, sub_txt)
    # return sub_txt


def findLineNumber(main_lst, search_string):
    """
    x = re.search("<!--Add The Box Here-->", txt)
    """
    Lines = []
    for i, each in enumerate(main_lst):
        if re.search(search_string, each):
            Lines.append(i)
            # yield i

    return Lines


def logger(line):
    with open('log.txt', 'a+') as f:
        f.write(str(line)+"\n")
"""
<!--Covid Test Predictions :<br>Test Results : Negative<br>Desription : XYZ<br>Comments : ABC<br><br> -->
"""

def rewriteHTMLTemplate(TestDetails,fpath):
    title = TestDetails["Title"]
    str_to_add = title + " - " +"<br>"
    del TestDetails["Title"]
    for each_key in TestDetails.keys():
        each_val = TestDetails[each_key]
        str_to_add = str_to_add + each_key + " : " + each_val +"<br>"
    txt = readHTML(fpath)
    linesadded = findLineNumber(txt, "<!--Add The Box Here-->")
    addtoHTMLFile(txt, pos=linesadded[-1]+1, str_to_add=str_to_add,fpath=fpath)
    return 1

if __name__ == "__main__":
    TemplateName = "AutoMail_Content_JINJA.html"
    TemplateFolder = "Templates"
    UserDetails = read_user_details()
    TemplatePath = os.path.join(ROOT_DIR,TemplateFolder, TemplateName)
    TempPath = os.path.join(ROOT_DIR,TemplateFolder, "temp.html")
    shutil.copy(TemplatePath,TempPath)
    string_res = """Covid Test Predictions :<br>Test Results : Negative<br>Desription : XYZ<br>Comments : ABC<br><br></p></td>"""
    TestDetails = {
                    "Title" : "Covid Test Predictions",
                    "Test Results" : "Positive",
                    "Desription" : "Level 4 - Advised Usage of Medical Treatment with Proper Consultation with a Doctor",
                    "Comments" : "Immediate Vaccination Needed."
                    }
    
    rewriteHTMLTemplate(TestDetails,fpath=TempPath)
    ReplacementDictionary = {'Your_Amazing_Name': 'Farhan Hai Khan',
                            'YOUR_AWESOME_NAME_HERE': 'FARHAN HAI KHAN',
                            'COVID': 'Coronavirus Scale : 91.53 %',
                            'Automated': 'Automated Tests : Failed',
                            'Manual': 'Manual Tests : Failed, Found COVID Positive',
                            'Description': 'Comments : Unfit for Travel',
                            'UID_Aadhaar': 'UID (Aadhaar) : 7458-2541-7364',
                            'Cert': 'Certificate ID : 9546-8741-9463',
                            'Generated': 'Generated : 19.04.2020 3:30GMT',
                            'Signature': 'Damik Dhar',
                            'Esteemed': 'DAMIK DHAR',
                            'Designation': 'Captain, Team MLXTREME',
                            'Autograph': 'Soumyadip Sarkar',
                            'Responsibility': 'SOUMYADIP SARKAR',
                            'Pos': 'CTO, MLXTREME'
                            }   
    InputFilepath =  os.path.join(MainDir,"main.docx")
    IntermediateDirectory='IntermediateFiles'
    create_dir(IntermediateDirectory)
    IntermediateFilepath = os.path.join(IntermediateDirectory,'new_main.docx')
    OutputDirec='Outputs'
    create_dir(OutputDirec)
    OutputFilepath = 'out.docx'
    OutputDirectory='/TemporaryFiles/'
    create_dir(OutputDirectory)
    """
    SrcImgFilePath='BarCodeImages/ReplacementImage.png'
    DestImgFileName='image4.png'
    """
        
    ReplaceText(InputFilepath, IntermediateFilepath, ReplacementDictionary)
    ExtractZipFile(IntermediateFilepath,OutputDirectory)
    """
    ReplaceWordImage(SrcImgFilePath,DestImgFileName,OutputDirectory)
    """
    ReconsructDocument(OutputDirectory,OutputFilepath)
    
    print('Docx File Written To :', OutputFilepath)
    
    MailDetails = {
                    "Subject": "Your Test Results via Doctor-in-a-Browser",
                    "To":      ['damikdhar@gmail.com','njrfarhandasilva10@gmail.com','nirmalya14misra@gmail.com','soumya997.sarkar@gmail.com'],
                    "Body" :   read_file(TempPath),
                    "AttachmentPath" : "out.docx"
                    }
    send_email_attach(MailDetails)
    print('Mail Sent!')
    # os.remove(TempPath)
    
    
