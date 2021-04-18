"""python PyCertGen.py"""
from docx import Document
import os
import pandas as pd
import copy
import common_utils


def print_list(lst):
    for each in lst:
        print(each)


def parser(document, v=1):
    """
    v : Verbosity
    """
    LINES = []
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if v == 1:
                print(text)
            LINES.append(text)
    return LINES


def cleanParsed(res):
    ans = []
    for each in res:
        stripped = each.strip()
        if not stripped == "":
            ans.append(stripped)
    return ans


def replacer(document, dic):
    newdocument = copy.deepcopy(document)  # deep copy
    for p in newdocument.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if text in dic.keys():
                text = text.replace(text, dic[text])
                inline[i].text = text
    return newdocument

# For various people load multiple dics


def DocxLoader(FileName):
    return Document(FileName)


def SingleSubstitution(InputDocxFilePath, OutputDocxFilePath, dic):
    document = Document(InputDocxFilePath)
    IsReplaced = replacer(document, dic)
    document.save(OutputDocxFilePath)


def CertGenEngine(InputDocxFilePath, InputExcelFilePath, OutputFolder):
    document = Document(InputDocxFilePath)
    df = pd.read_excel(InputExcelFilePath, engine="openpyxl")
    Attributes = list(df.columns)
    common_utils.create_dir(OutputFolder)
    for index, row in df.iterrows():
        dic = dict(row)
        newdocument = replacer(document, dic)
        OutputFileName = str(row[Attributes[0]]) + '.docx'
        # choose the output file names as first column name
        OutputFilePath = os.path.join(OutputFolder, OutputFileName)
        newdocument.save(OutputFilePath)


if __name__ == '__main__':
    ROOT_DIR = os.getcwd()
    FileName = os.path.join(ROOT_DIR, "CertTemplateSamples",
                            "Certificate_of_Appreciation.docx")
    document = Document(FileName)
    olddocument = document

    dic = {'YOUR_NAME': 'Farhan Hai Khan',
           'Outstanding_Professional_Experience': 'Machine Learning Engineer',
           'Date_Time': '12th of March 2021',
           'Sign': 'M. Agarwal',
           'Signatory_Name': 'Manish Agarwal',
           'Signatory_Position': 'IT Team Head',
           }
    Text_List = parser(document, v=0)
    Text_List_Parsed = cleanParsed(Text_List)
    print_list(Text_List_Parsed)
    newdocument = replacer(olddocument, dic)
    OutFileName = os.path.join(ROOT_DIR, "TempSaved", "Cert_Farhan.docx")
    newdocument.save(OutFileName)

    # Construct Multiple Docx Examples
    ExcelFileName = os.path.join(ROOT_DIR, "CertTemplateSamples",
                                 "Data_to_Fill.xlsx")
    df = pd.read_excel(ExcelFileName, engine="openpyxl")
    print(df.head())

    print(df.columns)

    Attributes = list(df.columns)

    for index, row in df.iterrows():
        print(dict(row))
    SaveFolder = os.path.join(ROOT_DIR, "TempSaved", "TempFiles")
    # create
    CertGenEngine(FileName, ExcelFileName, SaveFolder)

    SaveFolderPath = os.path.join("TempSaved", "TempFiles")
    # SaveZipFilePath = os.path.join(ROOT_DIR, "TempSaved", "CertGen.zip")
    SaveZipFilePath = "CertGen.zip"
    common_utils.zipper(SaveFolderPath, SaveZipFilePath)
